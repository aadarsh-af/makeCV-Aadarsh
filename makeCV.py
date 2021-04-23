from docx import Document
from docx.shared import Inches
import pyttsx3
from docx2pdf import convert


# make it speak
def speak(text):
    engine = pyttsx3.init()
    engine.setProperty('rate',160)
    voices = engine.getProperty('voices')
    engine.setProperty('voice',voices[1].id)
    pyttsx3.speak(text)


document = Document()
speak("Hello there! This is Adarsh's assistant and I'm here to help you make a great CV!")


# name
speak("So first of all, tell me your good name!")
name = input('\nWhat is your name? ')

#profile pic
speak(f'Hello {name}, It would be advisable if you make a CV which has your picture in it')
imgpermit = input("\nDo you want to add your photo in your CV? (y/n) ")
if imgpermit == 'y':
    # profile picture 
    path = input("Enter your image path: ")
    profile_pic = document.add_picture(path,width=Inches(2.0))
    profile_pic.alignment = 1

else:
    pass

#phone_number
speak("That's okay! Now I'd need you to enter your phone number?")
phone_number = input('\nEnter phone number: ')

#email
speak('Please Enter your email')
email = input('\nEnter email: ')

namedetails = document.add_paragraph(f'{name}\n{phone_number}\n{email}')

# about me
document.add_heading('About me')
speak("Okay it's going great! Now tell me about yourself like what jobs do you wish to apply for, etcetra")
document.add_paragraph(
    input('\nTell about yourself: ')
)

# work experience
speak("Have you ever worked in a company before?")
worked = input("\nHave you worked in a company before? (y/n) ")
if worked == 'y':
    speak("Please enter details")
    document.add_heading('Work Experience')
    p = document.add_paragraph()

    company = input('\nEnter Company Name: ')
    from_date = input('From Date: ')
    to_date = input('To Date: ')

    p.add_run(f'{company} ').bold = True
    p.add_run(f'({from_date} - {to_date}) \n').italic = True

    experience_details = input(f'Decribe your role at {company}: ')
    p.add_run(experience_details)

    # more experiences
    while True:
        has_more_experiences = input('\nDo you have more work experiences? (y/n) ')
        if has_more_experiences.lower() == 'y':
            p = document.add_paragraph()

            company = input('Enter company Name: ')
            from_date = input('From Date: ')
            to_date = input('To Date: ')

            p.add_run(f'{company} ').bold = True
            p.add_run(f'({from_date} - {to_date}) \n').italic = True

            experience_details = input(f'Decribe your role at {company}: ')
            p.add_run(experience_details)
        else:
            break

else:
    document.add_heading('Work Experience')
    p = document.add_paragraph()
    p.add_run('Fresher ').bold = True
    p.add_run('(Open to opportunities)').italic = True


# skills
speak("Besides your work experience, What skills do you have?")
document.add_heading('Skills')
skill = input('\nEnter a skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

speak("Any more of your skills you wanna tell me about?")
while True:
    has_more_skills = input('\nDo you have more skills? (y/n) ')
    if has_more_skills.lower() == 'y':
        skill = input('Enter another skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

speak(name + ", You have been so great")
speak("Now I will save your CV as a PDF file in this folder. Glad that I could help! Bye!")
# print("Thanks for your support.\nYour CV has been saved in this folder you can go save it as pdf and apply anywhere now!")

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "This CV is generated using Aadarsh Lalchandani's Make-CV project"
p.alignment = 1

document.save('My_CV.docx')

convert("My_CV.docx")
convert('docx/My_CV.docx','docx/My_CV.pdf')
convert("docx/")